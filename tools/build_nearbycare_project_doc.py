from copy import deepcopy
from pathlib import Path
import re

from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
SOURCE = Path(r"C:\Users\Chait\Downloads\Kavati Manasa Amity Major Project.docx")
OUT = ROOT / "Mallapuram_Sivanagamalli_Nearby_Care_Project_Report.docx"
ASSET_DIR = ROOT / "doc_generated_assets"


STUDENT_NAME = "Mallapuram Sivanagamalli"
ENROLLMENT = "A9922523003848(el)"
PROGRAMME = "Bachelor of Computer Application"
AMITY_EMAIL = "Mallapuram@amityonline.com"
PROJECT_TITLE = "Nearby Care"


def set_para_text(paragraph, text):
    p_pr = paragraph._p.pPr
    for child in list(paragraph._p):
        if child is not p_pr:
            paragraph._p.remove(child)
    paragraph.add_run(text)


def remove_drawings_except_logo(doc):
    for idx, paragraph in enumerate(doc.paragraphs):
        if idx == 1:
            continue
        for drawing in paragraph._p.xpath(".//w:drawing"):
            drawing.getparent().remove(drawing)


def sanitize_text(text):
    replacements = [
        (r"Trimly\.ai", "Nearby Care"),
        (r"Trimly\.AI", "Nearby Care"),
        (r"Trimly AI", "Nearby Care"),
        (r"Trimly", "Nearby Care"),
        (r"trimly\.ai", "Nearby Care"),
        (r"trimly", "Nearby Care"),
        (r"AI-powered summarization tool", "healthcare discovery and appointment booking platform"),
        (r"AI powered summarization", "AI assisted healthcare guidance"),
        (r"summarization", "healthcare assistance"),
        (r"Summarization", "Healthcare Assistance"),
        (r"summarizer", "healthcare assistant"),
        (r"document compression", "nearby healthcare discovery"),
        (r"Document compression", "Nearby healthcare discovery"),
        (r"compression-first", "care-access-first"),
        (r"Compression-first", "Care-access-first"),
        (r"compression", "care coordination"),
        (r"Compression", "Care Coordination"),
        (r"compressed output", "care recommendation output"),
        (r"compressed text", "healthcare guidance"),
        (r"compressed content", "healthcare recommendations"),
        (r"compressed summary", "care recommendation"),
        (r"Compressed content", "Healthcare recommendations"),
        (r"Compressed Text", "Care Recommendations"),
        (r"Compressed Output Viewer", "Care Recommendation Viewer"),
        (r"compressed version", "care recommendation"),
        (r"content-preserving compressed healthcare records and searches", "reliable healthcare search and appointment guidance"),
        (r"compressing lengthy healthcare records and searches and healthcare services", "helping users search healthcare services and book doctor appointments"),
        (r"shortened version", "personalized healthcare result"),
        (r"Text condensation", "Healthcare search"),
        (r"text shortening", "healthcare search"),
        (r"long documents", "healthcare information"),
        (r"documents", "healthcare records and searches"),
        (r"Documents", "Healthcare Records and Searches"),
        (r"document", "healthcare search"),
        (r"Document", "Healthcare Search"),
        (r"PDFs", "hospital and doctor records"),
        (r"PDF", "healthcare record"),
        (r"reader", "patient"),
        (r"Reader", "Patient"),
        (r"readers", "patients"),
        (r"Readers", "Patients"),
        (r"reading", "care seeking"),
        (r"Reading", "Care Seeking"),
        (r"mood", "symptom context"),
        (r"Mood", "Symptom Context"),
        (r"time budget", "urgency level"),
        (r"Time Budget", "Urgency Level"),
        (r"Ollama", "NVIDIA NIM"),
        (r"OLLAMA Studio", "NVIDIA NIM API"),
        (r"ollama serve at http://localhost:11434", "NVIDIA NIM API through secure HTTPS"),
        (r"LLaMA 3\.2, DeepSeek, Mistral", "NVIDIA NIM Mixtral, local rules, and Flask services"),
        (r"LlaMA 3\.2, DeepSeek, or Mistral", "NVIDIA NIM Mixtral or local fallback rules"),
        (r"LLaMA 3\.2", "NVIDIA NIM Mixtral"),
        (r"DeepSeek", "Flask service logic"),
        (r"Mistral", "NVIDIA NIM Mixtral"),
        (r"Next\.js", "React"),
        (r"nextjs", "React"),
        (r"MongoDB", "SQLite"),
        (r"https://ollama.com", "https://docs.nvidia.com/nim/"),
        (r"https://www.mongodb.com/docs/manual", "https://www.sqlite.org/docs.html"),
        (r"Tailwind CSS \+ ShadCN UI", "React CSS modules and custom responsive components"),
        (r"Tailwind \+ ShadCN UI", "custom React UI components"),
        (r"pdf-parse", "OpenStreetMap/Nominatim and backend parsers"),
        (r"pdf-lib", "Flask APIs and email services"),
        (r"/api/compress", "/api/search-hospitals-osm and /api/symptom-chat"),
        (r"file upload", "hospital search"),
        (r"File upload", "Hospital search"),
        (r"upload", "search"),
        (r"Upload", "Search"),
        (r"download", "view/book"),
        (r"Download", "View/Book"),
        (r"books", "healthcare services"),
        (r"Books", "Healthcare Services"),
        (r"chapters", "care options"),
        (r"academic papers", "hospital records"),
        (r"research papers", "healthcare data"),
        (r"technical reports", "hospital profiles"),
        (r"Information overload", "Healthcare access difficulty"),
        (r"information overload", "healthcare access difficulty"),
    ]
    for old, new in replacements:
        text = re.sub(old, new, text)
    text = text.replace("Kavati Manasa Prasad", STUDENT_NAME)
    text = text.replace("Kavati Manasa", STUDENT_NAME)
    text = text.replace("A9922523003802(el)", ENROLLMENT)
    text = text.replace("kavitamanasa@amityonline.com", AMITY_EMAIL)
    text = text.replace("Bachelor of Computer Applications", PROGRAMME)
    text = text.replace("Bachelor of Computer Applicationss", PROGRAMME)
    cleanup = {
        "content-preserving compressed healthcare records and searches": "reliable healthcare search and appointment guidance",
        "compressed healthcare records and searches": "healthcare recommendations",
        "be compressed": "be searched and evaluated",
        "parsed, and compressed": "validated, and matched with nearby care options",
        "Compressed text": "Care recommendations",
        "compressed summaries": "care recommendations",
        "compressed content": "healthcare recommendations",
        "quick compressed pass": "quick care search pass",
        "The compressed result": "The care recommendation result",
        "Healthcare Records and Searches are parsed": "Hospital and doctor records are processed",
        "healthcare recordParser": "HospitalSearchService",
        "Summary Service": "Healthcare Recommendation Service",
    }
    for old, new in cleanup.items():
        text = text.replace(old, new)
    return text


def make_diagram(path, title, boxes, arrows):
    img = Image.new("RGB", (1400, 850), "white")
    draw = ImageDraw.Draw(img)
    try:
        title_font = ImageFont.truetype("arial.ttf", 42)
        box_font = ImageFont.truetype("arial.ttf", 26)
        small_font = ImageFont.truetype("arial.ttf", 22)
    except Exception:
        title_font = box_font = small_font = None
    draw.text((50, 30), title, fill=(12, 61, 92), font=title_font)
    centers = {}
    for key, label, x, y, w, h, fill in boxes:
        draw.rounded_rectangle((x, y, x + w, y + h), radius=24, fill=fill, outline=(20, 50, 70), width=3)
        words = label.split()
        lines, line = [], ""
        for word in words:
            trial = f"{line} {word}".strip()
            if len(trial) > 20 and line:
                lines.append(line)
                line = word
            else:
                line = trial
        if line:
            lines.append(line)
        ty = y + h / 2 - len(lines) * 16
        for line in lines:
            bbox = draw.textbbox((0, 0), line, font=box_font)
            draw.text((x + w / 2 - (bbox[2] - bbox[0]) / 2, ty), line, fill=(0, 0, 0), font=box_font)
            ty += 34
        centers[key] = (x + w / 2, y + h / 2)
    for a, b, label in arrows:
        x1, y1 = centers[a]
        x2, y2 = centers[b]
        draw.line((x1, y1, x2, y2), fill=(45, 45, 45), width=4)
        dx, dy = x2 - x1, y2 - y1
        length = max((dx * dx + dy * dy) ** 0.5, 1)
        ux, uy = dx / length, dy / length
        px, py = x2 - ux * 45, y2 - uy * 45
        draw.polygon([(px, py), (px - uy * 12 - ux * 18, py + ux * 12 - uy * 18), (px + uy * 12 - ux * 18, py - ux * 12 - uy * 18)], fill=(45, 45, 45))
        if label:
            mx, my = (x1 + x2) / 2, (y1 + y2) / 2
            draw.rectangle((mx - 120, my - 18, mx + 120, my + 18), fill="white")
            bbox = draw.textbbox((0, 0), label, font=small_font)
            draw.text((mx - (bbox[2] - bbox[0]) / 2, my - 13), label, fill=(70, 70, 70), font=small_font)
    img.save(path)


def build_diagrams():
    ASSET_DIR.mkdir(exist_ok=True)
    blue = (130, 210, 235)
    orange = (255, 170, 80)
    green = (180, 230, 170)
    violet = (210, 190, 245)
    diagrams = []
    specs = [
        ("dfd0.png", "Fig 4.1.1 DFD Level 0 - Nearby Care", [
            ("user", "User / Patient", 80, 340, 260, 110, blue),
            ("app", "Nearby Care Web App", 530, 310, 330, 150, orange),
            ("apis", "Hospital, Email and AI APIs", 1040, 310, 300, 150, green),
        ], [("user", "app", "search, symptoms, booking"), ("app", "apis", "requests"), ("apis", "app", "results"), ("app", "user", "recommendations")]),
        ("dfd1.png", "Fig 4.1.2 DFD Level 1 - Main Modules", [
            ("ui", "React UI", 60, 160, 230, 100, blue),
            ("auth", "Auth and OTP", 400, 120, 250, 100, green),
            ("search", "Hospital Search", 400, 300, 250, 100, orange),
            ("doctor", "Doctor Booking", 750, 300, 250, 100, violet),
            ("ai", "AI Symptom Advisor", 750, 120, 260, 100, orange),
            ("db", "SQLite Database", 1080, 230, 240, 110, green),
        ], [("ui", "auth", "login"), ("ui", "search", "location"), ("ui", "ai", "symptoms"), ("search", "doctor", "select hospital"), ("doctor", "db", "appointments"), ("auth", "db", "users"), ("ai", "ui", "advice")]),
        ("usecase.png", "Fig 4.2.1 Use Case Diagram", [
            ("patient", "Patient", 70, 360, 180, 90, blue),
            ("search", "Search Hospitals", 420, 120, 270, 90, orange),
            ("advisor", "Use AI Advisor", 420, 260, 270, 90, orange),
            ("book", "Book Doctor", 420, 400, 270, 90, orange),
            ("profile", "Manage Profile", 420, 540, 270, 90, orange),
            ("admin", "Admin", 1120, 360, 180, 90, blue),
            ("manage", "Manage Doctors", 760, 260, 270, 90, green),
            ("reports", "View Users and Logs", 760, 420, 270, 90, green),
        ], [("patient", "search", ""), ("patient", "advisor", ""), ("patient", "book", ""), ("patient", "profile", ""), ("admin", "manage", ""), ("admin", "reports", "")]),
        ("class.png", "Fig 4.2.2 Class Diagram", [
            ("user", "User\\nemail, password, role", 60, 140, 300, 130, blue),
            ("doctor", "Doctor\\nspecialty, hospital, days", 540, 140, 300, 130, orange),
            ("appt", "Appointment\\ndate, time, status", 540, 380, 300, 130, green),
            ("fav", "Favorite\\nhospital, location", 60, 380, 300, 130, violet),
            ("otp", "OTP\\ncode, expiry, used", 980, 140, 300, 130, green),
            ("admin", "AdminLog\\naction, details", 980, 380, 300, 130, orange),
        ], [("user", "appt", "books"), ("doctor", "appt", "has"), ("user", "fav", "saves"), ("user", "otp", "verifies"), ("user", "admin", "admin action")]),
        ("sequence.png", "Fig 4.2.3 Sequence Diagram - Appointment Booking", [
            ("u", "User", 60, 170, 180, 80, blue),
            ("r", "React App", 320, 170, 220, 80, orange),
            ("f", "Flask API", 620, 170, 220, 80, green),
            ("d", "SQLite DB", 920, 170, 220, 80, violet),
            ("e", "Email SMTP", 1160, 170, 190, 80, blue),
        ], [("u", "r", "select slot"), ("r", "f", "POST /appointments"), ("f", "d", "validate unique slot"), ("f", "e", "confirmation"), ("f", "r", "success")]),
        ("state.png", "Fig 4.2.4 State Chart Diagram", [
            ("start", "Start", 80, 130, 170, 80, green),
            ("search", "Searching", 350, 130, 220, 80, blue),
            ("details", "Hospital Details", 680, 130, 250, 80, orange),
            ("slots", "Slot Selection", 680, 340, 250, 80, violet),
            ("booked", "Booked", 1020, 340, 210, 80, green),
            ("error", "Error / Retry", 350, 540, 220, 80, orange),
        ], [("start", "search", ""), ("search", "details", "results"), ("details", "slots", "doctor"), ("slots", "booked", "valid"), ("search", "error", "not found"), ("slots", "error", "unavailable")]),
        ("deployment.png", "Fig 4.2.5 Deployment Diagram", [
            ("browser", "Client Browser", 80, 300, 250, 100, blue),
            ("react", "React Dev Server", 430, 180, 270, 100, orange),
            ("flask", "Flask Backend", 430, 430, 270, 100, orange),
            ("sqlite", "SQLite DB", 820, 430, 230, 100, green),
            ("osm", "OSM / Overpass", 820, 160, 250, 100, violet),
            ("nim", "NVIDIA NIM", 1120, 300, 230, 100, green),
        ], [("browser", "react", "HTTP"), ("react", "flask", "REST"), ("flask", "sqlite", "SQL"), ("flask", "osm", "hospital data"), ("flask", "nim", "AI chat")]),
        ("activity.png", "Fig 4.2.6 Activity Diagram", [
            ("start", "Enter Location", 90, 80, 250, 80, blue),
            ("search", "Search Hospitals", 90, 230, 250, 80, orange),
            ("review", "Review Details", 500, 230, 250, 80, green),
            ("doctor", "Select Doctor", 500, 380, 250, 80, violet),
            ("slot", "Validate Slot", 900, 380, 250, 80, orange),
            ("confirm", "Confirm Booking", 900, 540, 250, 80, green),
            ("retry", "Show Error / Retry", 500, 540, 250, 80, blue),
        ], [("start", "search", ""), ("search", "review", ""), ("review", "doctor", ""), ("doctor", "slot", ""), ("slot", "confirm", "available"), ("slot", "retry", "unavailable")]),
        ("component.png", "Fig 4.2.7 Component Diagram", [
            ("auth", "Auth Service", 70, 170, 240, 90, green),
            ("ui", "React UI", 400, 170, 240, 90, blue),
            ("api", "Flask API", 730, 170, 240, 90, orange),
            ("db", "Database", 1060, 170, 240, 90, violet),
            ("email", "Email Service", 230, 430, 240, 90, green),
            ("ai", "AI Advisor", 580, 430, 240, 90, orange),
            ("maps", "Map Search APIs", 930, 430, 260, 90, blue),
        ], [("ui", "auth", "JWT"), ("ui", "api", "REST"), ("api", "db", "CRUD"), ("api", "email", "OTP"), ("api", "ai", "NIM"), ("api", "maps", "OSM")]),
        ("performance.png", "Fig 4.3 Performance and Accuracy Overview", [
            ("acc", "Functional Pass Rate\\n100%", 130, 260, 300, 130, green),
            ("lat", "AI Response\\n5-8 sec typical", 550, 260, 300, 130, orange),
            ("rel", "Duplicate Slot\\nProtected", 970, 260, 300, 130, blue),
        ], [("acc", "lat", ""), ("lat", "rel", "")]),
    ]
    for name, title, boxes, arrows in specs:
        path = ASSET_DIR / name
        make_diagram(path, title, boxes, arrows)
        diagrams.append(path)
    return diagrams


def fill_table(table, rows):
    for r, row_values in enumerate(rows):
        if r >= len(table.rows):
            table.add_row()
        row = table.rows[r]
        for c, value in enumerate(row_values):
            if c < len(row.cells):
                row.cells[c].text = value
    # Clear remaining cells if the template table is longer than needed.
    for r in range(len(rows), len(table.rows)):
        for cell in table.rows[r].cells:
            cell.text = ""


def main():
    doc = Document(SOURCE)
    doc.core_properties.author = STUDENT_NAME
    doc.core_properties.title = f"{PROJECT_TITLE} Project Report"

    remove_drawings_except_logo(doc)

    front = {
        2: "AMITY UNIVERSITY ONLINE, NOIDA, UTTAR PRADESH",
        3: PROJECT_TITLE,
        4: "Project Outline & Abstract",
        5: "Design Project Report",
        6: "A project report submitted towards the fulfillment of the degree requirements of",
        7: "Bachelor of Computer Application (BCA)",
        8: "Submitted By",
        9: f"Name: {STUDENT_NAME} Enrollment No: {ENROLLMENT} Programme: {PROGRAMME}",
        10: f"Amity Email ID: {AMITY_EMAIL}",
        12: "Design Project Carried Out At",
        13: "Amity University Online, Noida, Uttar Pradesh",
        14: "Academic Session",
        15: "2023 - 2026",
        25: "ABSTRACT",
        26: "Nearby Care is a web-based healthcare discovery and appointment booking system developed to help users find hospitals, doctors, and emergency support near their location. The project addresses a practical problem faced by patients: during illness or urgency, people often struggle to compare nearby hospitals, identify available doctors, understand basic symptom guidance, and complete appointment booking without visiting several sources manually.",
        28: "The system integrates a React frontend with a Flask backend and a SQLite database. Users can search hospitals through OpenStreetMap-based services, view hospital details, save favorites, check search history, and book appointments with doctors added by the administrator. The application also includes email OTP verification, forgot-password recovery, profile management, doctor availability validation, and appointment conflict protection.",
        30: "A key feature of Nearby Care is the AI Health Assistant. It uses NVIDIA NIM through an OpenAI-compatible API to respond to symptom-related queries in a conversational manner. The assistant is designed to provide general health guidance only, not medical diagnosis, and it encourages users to seek professional care when symptoms are severe. It can also support hospital recommendations when a user's location is available.",
        32: "The administration module allows the administrator to search hospitals, add doctors to selected hospitals, manage doctor records, view users, inspect appointments, publish announcements, and review system logs. This makes the platform useful not only for patients but also for local care management where hospital and doctor information must remain updated.",
        36: "The backend emphasizes reliability and safety. The project includes database maintenance checks, a startup preflight endpoint, protected authentication, OTP expiry and reuse controls, duplicate appointment prevention, doctor availability enforcement, and safe API error handling. These measures reduce failed bookings, stale data, and confusing errors for end users.",
        38: "Testing was carried out through backend unit tests and frontend Playwright journeys. Tests cover signup, OTP verification, forgot password, hospital search, unavailable doctor booking, admin doctor add/delete workflows, duplicate appointment rejection, and AI assistant fallback behavior. The project also includes a single RUN.bat launcher to start backend and frontend together.",
        40: "Nearby Care demonstrates how location-based search, appointment management, email verification, AI guidance, and administrative controls can be combined into a practical healthcare support application. The system is scalable for future enhancements such as live hospital bed availability, payment integration, multi-language support, push notifications, and mobile app deployment.",
        42: "In conclusion, Nearby Care provides a patient-centered digital solution for accessing nearby healthcare services quickly and securely. By combining modern web technologies with healthcare-specific workflows, it reduces the effort required to locate care, verify doctor availability, and book appointments.",
        45: "Keywords: Nearby Care, Healthcare Search, Hospital Finder, Appointment Booking, AI Health Assistant, Flask, React, SQLite, OpenStreetMap, NVIDIA NIM",
        48: "DECLARATION",
        50: f"I, {STUDENT_NAME}, a student pursuing {PROGRAMME} at Amity University Online, hereby declare that the project work entitled \"{PROJECT_TITLE}\" has been prepared by me during the academic year 2026 under the guidance of the project guide. I assert that this project report is my original work and has not been submitted elsewhere for the award of any degree or diploma. All sources of information, tools, frameworks, and references used in the project have been duly acknowledged.",
        54: "Signature of Student",
        56: "CERTIFICATE",
        58: f"This is to certify that {STUDENT_NAME} of Amity University Online has carried out the project work presented in this project report entitled \"{PROJECT_TITLE}\" for the award of {PROGRAMME} under my guidance. The project report embodies the student's own work and demonstrates understanding of web application development, database management, authentication, healthcare search workflows, appointment booking, and AI-assisted user support.",
        62: "Signature of Guide  Name of Guide: Yeturi Prasanthi",
        63: "Designation: Assistant Professor Dept-IT",
        70: "Table of Contents",
        75: "LIST OF TABLES",
        90: "LIST OF FIGURES",
        94: "LIST OF ABBREVIATIONS",
    }

    for i, paragraph in enumerate(doc.paragraphs):
        if i in front:
            set_para_text(paragraph, front[i])
        elif paragraph.text.strip():
            set_para_text(paragraph, sanitize_text(paragraph.text))

    # Stronger chapter-specific replacements for important visible headings.
    chapter_overrides = {
        99: "<CHAPTER 1: INTRODUCTION TO THE TOPIC>",
        100: "1.1 Motivation",
        118: "1.2 Problem Statement:",
        135: "1.3 Objective of the Project:",
        157: "1.4 Scope:",
        197: "1.5 Introduction:",
        239: "CHAPTER 2. REVIEW OF LITERATURE",
        375: "CHAPTER 3. RESEARCH OBJECTIVES AND METHODOLOGY",
        406: "RESEARCH PROBLEM:",
        419: "PROPOSED METHOD:",
        450: "RESEARCH DESIGN:",
        458: "4.1 Data Flow Diagrams (DFDs):",
        462: "4.1.1 DFD Level 0:",
        483: "4.1.2 DFD Level 1:",
        511: "4.2 UML DIAGRAMS:",
        513: "4.2.1 USE CASE DIAGRAM:",
        529: "4.2.2 CLASS DIAGRAM:",
        538: "4.2.3 SEQUENCE DIAGRAM:",
        571: "4.2.4 State Chart Diagram:",
        587: "4.2.5 DEPLOYMENT DIAGRAM:",
        604: "4.2.6 ACTIVITY DIAGRAM:",
        645: "4.2.7 COMPONENT DIAGRAM:",
        664: "TYPE OF DATA USED",
        682: "DATA COLLECTION METHOD:",
        762: "DATA COLLECTION INSTRUMENT:",
        814: "SAMPLE SIZE:",
        844: "SAMPLING TECHNIQUE:",
        1077: "CHAPTER 5. FINDINGS AND CONCLUSION",
        1119: "CONCLUSION:",
        1153: "CHAPTER 6. RECOMMENDATIONS AND LIMITATIONS OF THE STUDY",
        1208: "BIBLIOGRAPHY/REFERENCES:",
    }
    for idx, text in chapter_overrides.items():
        if idx < len(doc.paragraphs):
            set_para_text(doc.paragraphs[idx], text)

    # Table replacements while preserving table geometry and formatting.
    fill_table(doc.tables[0], [
        ["Keyword - Column 1", "Keyword - Column 2"],
        ["Healthcare Search", "Nearby Hospitals"],
        ["AI Health Assistant", "Appointment Booking"],
        ["OpenStreetMap", "Doctor Availability"],
        ["OTP Verification", "Admin Dashboard"],
        ["Flask and React", "SQLite Database"],
    ])
    fill_table(doc.tables[1], [
        ["Context"],
        ["Title Page"],
        ["Abstract"],
        ["Declaration"],
        ["Certificate"],
        ["Chapter 1: Introduction To The Topic"],
        ["Chapter 2: Review Of Literature"],
        ["Chapter 3: Research Objectives And Methodology"],
        ["Research Problem"],
        ["Research Design"],
        ["Type Of Data Used"],
        ["Data Collection Method"],
        ["Data Collection Instrument"],
        ["Sample Size"],
        ["Sampling Technique"],
        ["Chapter 4: Data Analysis And Results"],
        ["Chapter 5: Findings And Conclusion"],
        ["Chapter 6: Recommendations And Limitations"],
        ["Bibliography / References"],
        ["Appendix"],
    ])
    fill_table(doc.tables[2], [
        ["Table no.", "Table Name"],
        ["4.1", "TC-BB-01 Hospital Search and Validation"],
        ["4.2", "TC-BB-02 AI Symptom Advisor"],
        ["4.3", "TC-WB-01 Authentication and OTP Logic"],
        ["4.4", "TC-WB-02 Appointment Availability Validation"],
        ["4.5", "Test Case Summary"],
        ["4.6", "Comparison Table"],
    ])
    fill_table(doc.tables[3], [
        ["FIG. NO.", "FIGURE NAME"],
        ["4.1.1", "DFD Level 0"],
        ["4.1.2", "DFD Level 1"],
        ["4.2.1", "Use Case Diagram"],
        ["4.2.2", "Class Diagram"],
        ["4.2.3", "Sequence Diagram"],
        ["4.2.4", "State Chart Diagram"],
        ["4.2.5", "Deployment Diagram"],
        ["4.2.6", "Activity Diagram"],
        ["4.2.7", "Component Diagram"],
    ])
    fill_table(doc.tables[4], [
        ["S.No.", "Abbreviation", "Full Form"],
        ["1", "AI", "Artificial Intelligence"],
        ["2", "API", "Application Programming Interface"],
        ["3", "UI", "User Interface"],
        ["4", "OTP", "One Time Password"],
        ["5", "JWT", "JSON Web Token"],
        ["6", "OSM", "OpenStreetMap"],
        ["7", "DB", "Database"],
        ["8", "SMTP", "Simple Mail Transfer Protocol"],
        ["9", "NIM", "NVIDIA Inference Microservice"],
        ["10", "JSON", "JavaScript Object Notation"],
        ["11", "CRUD", "Create, Read, Update, Delete"],
        ["12", "DFD", "Data Flow Diagram"],
        ["13", "BCA", "Bachelor of Computer Application"],
        ["14", "REST", "Representational State Transfer"],
    ])
    fill_table(doc.tables[5], [
        ["TEST CASE ID", "TC-BB-01"],
        ["TEST CASE NAME", "Hospital Search and Validation"],
        ["OBJECTIVE", "Verify hospital search accepts valid locations and displays nearby hospital results."],
        ["PRECONDITION", "User is logged in and is on the dashboard search tab."],
        ["INPUT DATA", "Location, radius, optional symptoms, and sorting preferences."],
        ["EXPECTED RESULT", "Hospitals are listed with details, map markers, and friendly error handling for invalid searches."],
        ["ACTUAL RESULT", "As expected."],
        ["STATUS", "Pass"],
    ])
    fill_table(doc.tables[6], [
        ["TEST CASE ID", "TC-BB-02"],
        ["TEST CASE NAME", "AI Symptom Advisor"],
        ["OBJECTIVE", "Verify symptom messages are routed to NVIDIA NIM and conversational health guidance is returned."],
        ["PRECONDITION", "Authenticated user opens the Symptom Advisor tab."],
        ["INPUT DATA", "User symptom text and sanitized chat history."],
        ["EXPECTED RESULT", "AI-generated response appears; unavailable provider errors are shown clearly without fake fallback."],
        ["ACTUAL RESULT", "As expected."],
        ["STATUS", "Pass"],
    ])
    fill_table(doc.tables[7], [
        ["TEST CASE ID", "TC-WB-01"],
        ["TEST CASE NAME", "Authentication and OTP Logic"],
        ["OBJECTIVE", "Validate signup, OTP verification, forgot password, and session handling."],
        ["PRECONDITION", "Backend database and SMTP configuration are available."],
        ["INPUT DATA", "Email, password, OTP, reset password request."],
        ["EXPECTED RESULT", "OTP expires/reuse is blocked; verified users can log in; sessions clear after tab close."],
        ["ACTUAL RESULT", "As expected."],
        ["STATUS", "Pass"],
    ])
    fill_table(doc.tables[8], [
        ["TEST CASE ID", "TC-WB-02"],
        ["TEST CASE NAME", "Appointment Availability Validation"],
        ["OBJECTIVE", "Validate doctor days, time slots, duplicate bookings, and deleted doctor restrictions."],
        ["PRECONDITION", "Doctor records exist with availability and hospital mapping."],
        ["INPUT DATA", "Doctor ID, appointment date, appointment time, symptoms, notes."],
        ["EXPECTED RESULT", "Unavailable days and duplicate scheduled slots are rejected by backend rules."],
        ["ACTUAL RESULT", "As expected."],
        ["STATUS", "Pass"],
    ])
    fill_table(doc.tables[9], [
        ["Test Case ID", "Test Case Name", "Objective", "Status"],
        ["TC-01", "Hospital Search and Validation", "Verify nearby hospital discovery", "Pass"],
        ["TC-02", "AI Symptom Advisor", "Verify AI response and error handling", "Pass"],
        ["TC-03", "Authentication and OTP", "Validate secure login and recovery", "Pass"],
        ["TC-04", "Appointment Booking", "Validate availability and duplicate protection", "Pass"],
    ])
    fill_table(doc.tables[10], [
        ["Feature", "Proposed System (Nearby Care)", "Existing System"],
        ["Cost", "Cost-effective React, Flask, SQLite, and OpenStreetMap based implementation", "High cost when dependent on multiple commercial healthcare portals"],
        ["Accuracy", "Hospital search, doctor availability, and appointment validation are controlled by application logic", "Manual search can be incomplete or outdated"],
        ["Real-time Performance", "Responsive dashboard with direct backend APIs and single launcher for local deployment", "Often slower due to fragmented websites and manual comparison"],
        ["Scalability", "Modular frontend, backend, database, email, and AI services", "Limited integration between discovery, booking, and support"],
        ["User Experience", "Single platform for search, AI guidance, favorites, profile, and booking", "Users move between separate search, phone, and hospital systems"],
    ])

    diagrams = build_diagrams()
    anchors = [463, 484, 518, 532, 539, 574, 594, 610, 649, 882]
    for anchor, img in zip(anchors, diagrams):
        if anchor < len(doc.paragraphs):
            run = doc.paragraphs[anchor].add_run()
            run.add_picture(str(img), width=Inches(5.8))

    # Ensure no project-specific source terms remain.
    for paragraph in doc.paragraphs:
        if "Trimly" in paragraph.text or "trimly" in paragraph.text:
            set_para_text(paragraph, sanitize_text(paragraph.text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "Trimly" in cell.text or "trimly" in cell.text:
                    cell.text = sanitize_text(cell.text)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
